from __future__ import annotations
import json, csv, re
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

EXPORT_DIR = Path('./data/exports')
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

def slugify(s: str) -> str:
    return re.sub(r'[^a-zA-Z0-9_-]+', '_', s)[:80].strip('_') or 'export'

def export_job(job: dict, fmt: str) -> Path:
    fmt = fmt.lower().strip('.')
    base = f"job_{job['id']}_{slugify(job['project_slug'])}"
    output = job.get('output') or ''
    if fmt == 'json':
        path = EXPORT_DIR / f'{base}.json'
        payload = dict(job)
        for k in ['inputs_json','customization_json','uploaded_file_ids_json','usage_json']:
            if k in payload and isinstance(payload[k], str):
                try:
                    payload[k] = json.loads(payload[k])
                except Exception:
                    pass
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        return path
    if fmt in {'md','markdown'}:
        path = EXPORT_DIR / f'{base}.md'
        path.write_text(output, encoding='utf-8')
        return path
    if fmt == 'txt':
        path = EXPORT_DIR / f'{base}.txt'
        path.write_text(output, encoding='utf-8')
        return path
    if fmt == 'csv':
        path = EXPORT_DIR / f'{base}.csv'
        with path.open('w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['job_id','project','provider','status','output'])
            writer.writerow([job['id'], job['project_slug'], job['provider_name'], job['status'], output])
        return path
    if fmt == 'docx':
        path = EXPORT_DIR / f'{base}.docx'
        doc = Document()
        doc.add_heading(f"{job['project_slug']} — Job {job['id']}", 0)
        for para in output.split('\n'):
            if para.startswith('# '):
                doc.add_heading(para[2:], 1)
            elif para.startswith('## '):
                doc.add_heading(para[3:], 2)
            elif para.strip():
                doc.add_paragraph(para)
        doc.save(path)
        return path
    if fmt == 'pdf':
        path = EXPORT_DIR / f'{base}.pdf'
        styles = getSampleStyleSheet()
        story = [Paragraph(f"{job['project_slug']} — Job {job['id']}", styles['Title']), Spacer(1, 12)]
        for para in output.split('\n'):
            if not para.strip():
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(para.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'), styles['BodyText']))
                story.append(Spacer(1, 5))
        SimpleDocTemplate(str(path), pagesize=A4).build(story)
        return path
    raise ValueError(f'Unsupported export format: {fmt}')
